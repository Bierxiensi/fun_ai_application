from docx import Document as DocxDocument

class WordExtractor():
    def extract(self, file_path):
        """解析 DOCX 文件，提取全部文本"""
        doc = DocxDocument(file_path)
        all_text = []

        # 提取段落文本
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                all_text.append(paragraph.text.strip())

        # 提取表格文本
        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    table_text.append(" | ".join(row_text))
            if table_text:
                all_text.append("\n".join(table_text))

        return "\n".join(all_text)